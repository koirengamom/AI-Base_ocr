"""
output_writer.py
-----------------
Takes raw OCR results and saves to both .txt and .docx files, preserving
table layout (rows/columns) where detected, instead of merging table
cells into jumbled paragraph text.
"""

import os
import re
from datetime import datetime

from table_detector import detect_blocks

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


_NUMBERING_MARKER_RE = re.compile(
    r"^\(?(\d{1,3}|[a-zA-Z]{1,3}|[ivxlcdm]{1,6})\)?[.)]?$"
)


def _is_bare_marker(text: str) -> bool:
    candidate = text.strip()
    if not candidate or len(candidate) > 6:
        return False
    return bool(_NUMBERING_MARKER_RE.match(candidate))


def _line_height(item: dict) -> float:
    ys = [point[1] for point in item["box"]]
    return max(ys) - min(ys)


def _line_top(item: dict) -> float:
    return min(point[1] for point in item["box"])


def _line_bottom(item: dict) -> float:
    return max(point[1] for point in item["box"])


def _safe_conf(value) -> float:
    return float(value) if isinstance(value, (int, float)) else 0.0


def merge_lines_into_paragraphs(results: list, gap_ratio: float = 0.6) -> list:
    if not results:
        return []

    paragraphs = []
    current_texts = [results[0]["text"]]
    current_confidences = [_safe_conf(results[0]["confidence"])]
    current_ref = results[0]

    for item in results[1:]:
        prev_is_marker = _is_bare_marker(current_texts[-1])

        gap = _line_top(item) - _line_bottom(current_ref)
        avg_height = (_line_height(current_ref) + _line_height(item)) / 2 or 1
        close_enough = gap <= gap_ratio * avg_height

        if prev_is_marker or close_enough:
            if prev_is_marker:
                current_texts[-1] = f"{current_texts[-1]} {item['text']}".strip()
            else:
                current_texts.append(item["text"])
            current_confidences.append(_safe_conf(item["confidence"]))
            current_ref = item
        else:
            paragraphs.append({
                "text": " ".join(current_texts).strip(),
                "confidence": round(sum(current_confidences) / len(current_confidences), 4),
            })
            current_texts = [item["text"]]
            current_confidences = [_safe_conf(item["confidence"])]
            current_ref = item

    paragraphs.append({
        "text": " ".join(current_texts).strip(),
        "confidence": round(sum(current_confidences) / len(current_confidences), 4),
    })

    return paragraphs


def format_results_for_display(results: list, show_confidence: bool = True) -> str:
    if not results:
        return "[No text detected]"
    lines = []
    for item in results:
        if show_confidence:
            conf_pct = _safe_conf(item["confidence"]) * 100
            lines.append(f"{item['text']}    [confidence: {conf_pct:.2f}%]")
        else:
            lines.append(item["text"])
    return "\n".join(lines)


def calculate_average_confidence(results: list) -> float:
    if not results:
        return 0.0
    total = sum(_safe_conf(item["confidence"]) for item in results)
    return round((total / len(results)) * 100, 2)


def _flatten_blocks(results):
    """Yields ('para', [ocr items]) or ('table', [[str,...],...]) in original order."""
    group = []
    for b in detect_blocks(results):
        if b["type"] == "table":
            if group:
                yield ("para", group)
                group = []
            yield ("table", b["rows"])
        else:
            group.extend(b["items"])
    if group:
        yield ("para", group)


def _render_table_md(rows):
    return "\n".join("| " + " | ".join(r) + " |" for r in rows)


def format_plain_text(results: list) -> str:
    if not results:
        return ""
    out = []
    for kind, payload in _flatten_blocks(results):
        if kind == "table":
            out.append(_render_table_md(payload))
        else:
            out.extend(item["text"] for item in merge_lines_into_paragraphs(payload))
    return "\n".join(out)


def _add_blocks_to_doc(doc, results):
    for kind, payload in _flatten_blocks(results):
        if kind == "table":
            ncols = max(len(r) for r in payload)
            table = doc.add_table(rows=0, cols=ncols)
            table.style = "Table Grid"
            for r in payload:
                cells = table.add_row().cells
                for idx in range(ncols):
                    cells[idx].text = r[idx] if idx < len(r) else ""
        else:
            for item in merge_lines_into_paragraphs(payload):
                doc.add_paragraph(item["text"])


def _build_image_docx(results: list, source_filename: str):
    doc = Document()
    _add_blocks_to_doc(doc, results)
    return doc


def _build_pdf_docx(page_results: list, source_filename: str):
    from docx.enum.text import WD_BREAK  # noqa: F401 (kept for compat)

    doc = Document()
    for page_num, results in enumerate(page_results, start=1):
        if results:
            _add_blocks_to_doc(doc, results)
        else:
            doc.add_paragraph("[No text detected on this page]")
        if page_num < len(page_results):
            doc.add_page_break()
    return doc


def save_image_ocr_output(results: list, source_filename: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(source_filename))[0]

    txt_path = os.path.join(output_dir, f"{base_name}_ocr_output.txt")
    avg_conf = calculate_average_confidence(results)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    header = (
        f"OCR Output Report\n"
        f"{'=' * 50}\n"
        f"Source File       : {os.path.basename(source_filename)}\n"
        f"Processed On      : {timestamp}\n"
        f"Lines Detected    : {len(results)}\n"
        f"Average Confidence: {avg_conf}%\n"
        f"{'=' * 50}\n\n"
    )
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(header)
        f.write(format_plain_text(results))
        f.write("\n")

    docx_path = None
    if DOCX_AVAILABLE:
        docx_path = os.path.join(output_dir, f"{base_name}_ocr_output.docx")
        _build_image_docx(results, source_filename).save(docx_path)

    return txt_path, docx_path


def save_pdf_ocr_output(page_results: list, source_filename: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(source_filename))[0]

    txt_path = os.path.join(output_dir, f"{base_name}_ocr_output.txt")
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    total_lines = sum(len(p) for p in page_results)
    all_conf = [_safe_conf(item["confidence"]) for page in page_results for item in page]
    avg_conf = round((sum(all_conf) / len(all_conf)) * 100, 2) if all_conf else 0.0
    header = (
        f"OCR Output Report (Multi-Page PDF)\n"
        f"{'=' * 50}\n"
        f"Source File       : {os.path.basename(source_filename)}\n"
        f"Processed On      : {timestamp}\n"
        f"Total Pages       : {len(page_results)}\n"
        f"Total Lines       : {total_lines}\n"
        f"Average Confidence: {avg_conf}%\n"
        f"{'=' * 50}\n\n"
    )
    body_parts = []
    for page_num, results in enumerate(page_results, start=1):
        page_avg = calculate_average_confidence(results)
        page_header = f"\n--- Page {page_num} (Lines: {len(results)}, Avg Confidence: {page_avg}%) ---\n"
        page_text = format_plain_text(results) if results else "[No text detected on this page]"
        body_parts.append(page_header + page_text)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(header)
        f.write("\n".join(body_parts))
        f.write("\n")

    docx_path = None
    if DOCX_AVAILABLE:
        docx_path = os.path.join(output_dir, f"{base_name}_ocr_output.docx")
        _build_pdf_docx(page_results, source_filename).save(docx_path)

    return txt_path, docx_path